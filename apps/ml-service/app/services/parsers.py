"""Document parsers for PDF, DOCX, and plain text files.

Extracts text content and converts to structured records suitable
for the cleaning pipeline and downstream training use.

Supports:
- PDF (via PyMuPDF / fitz)
- DOCX (via python-docx)
- TXT / Markdown (plain text)

Each parser returns a list of records, where each record represents
a logical section (page, paragraph, or chunk) of the document.
"""

import logging
import os
import re
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_file(filepath: str, *, chunk_size: int = 1000) -> list[dict]:
    """Auto-detect file type and parse to records.

    Args:
        filepath: Path to the file to parse.
        chunk_size: Maximum characters per chunk for plain text files.

    Returns:
        List of dicts, each with 'text', 'source', and 'metadata' keys.
    """
    ext = Path(filepath).suffix.lower()
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_docx,
        ".txt": parse_text,
        ".md": parse_text,
        ".markdown": parse_text,
    }

    parser = parsers.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(parsers.keys())}")

    return parser(filepath, chunk_size=chunk_size)


def parse_pdf(filepath: str, *, chunk_size: int = 1000) -> list[dict]:
    """Extract text from a PDF file, one record per page.

    Uses PyMuPDF (fitz) for fast, accurate text extraction including
    table-like structures and multi-column layouts.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required for PDF parsing. Install with: pip install PyMuPDF"
        ) from None

    filename = os.path.basename(filepath)
    records: list[dict] = []

    doc = fitz.open(filepath)
    try:
        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text").strip()

            if not text:
                continue

            # Clean up common PDF artifacts
            text = _clean_pdf_text(text)

            records.append(
                {
                    "text": text,
                    "source": filename,
                    "metadata": {
                        "page": page_num,
                        "total_pages": len(doc),
                        "char_count": len(text),
                        "format": "pdf",
                    },
                }
            )
    finally:
        doc.close()

    logger.info(
        "[Parser] PDF '%s': %d pages → %d records",
        filename,
        len(doc),
        len(records),
    )
    return records


def parse_docx(filepath: str, *, chunk_size: int = 1000) -> list[dict]:
    """Extract text from a DOCX file, one record per paragraph or heading group.

    Groups consecutive paragraphs under the same heading into a single record.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install python-docx"
        ) from None

    filename = os.path.basename(filepath)
    doc = Document(filepath)
    records: list[dict] = []

    current_heading = "Document"
    current_text_parts: list[str] = []
    section_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this paragraph is a heading
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")

        if is_heading:
            # Flush previous section
            if current_text_parts:
                section_text = "\n".join(current_text_parts)
                records.append(
                    {
                        "text": section_text,
                        "source": filename,
                        "metadata": {
                            "heading": current_heading,
                            "section": section_idx,
                            "char_count": len(section_text),
                            "format": "docx",
                        },
                    }
                )
                section_idx += 1
                current_text_parts = []

            current_heading = text
        else:
            current_text_parts.append(text)

    # Flush remaining content
    if current_text_parts:
        section_text = "\n".join(current_text_parts)
        records.append(
            {
                "text": section_text,
                "source": filename,
                "metadata": {
                    "heading": current_heading,
                    "section": section_idx,
                    "char_count": len(section_text),
                    "format": "docx",
                },
            }
        )

    # Also extract tables
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        if rows:
            table_text = "\n".join(rows)
            records.append(
                {
                    "text": table_text,
                    "source": filename,
                    "metadata": {
                        "type": "table",
                        "table_index": table_idx,
                        "row_count": len(rows),
                        "char_count": len(table_text),
                        "format": "docx",
                    },
                }
            )

    logger.info(
        "[Parser] DOCX '%s': %d paragraphs → %d records",
        filename,
        len(doc.paragraphs),
        len(records),
    )
    return records


def parse_text(filepath: str, *, chunk_size: int = 1000) -> list[dict]:
    """Parse a plain text or markdown file into chunked records.

    Splits on double newlines (paragraphs), then groups into chunks
    up to `chunk_size` characters.
    """
    filename = os.path.basename(filepath)

    with open(filepath, encoding="utf-8", errors="replace") as f:
        content = f.read()

    # Split into paragraphs
    paragraphs = re.split(r"\n\s*\n", content)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    # Group paragraphs into chunks
    records: list[dict] = []
    current_chunk: list[str] = []
    current_length = 0
    chunk_idx = 0

    for para in paragraphs:
        if current_length + len(para) > chunk_size and current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            records.append(
                {
                    "text": chunk_text,
                    "source": filename,
                    "metadata": {
                        "chunk": chunk_idx,
                        "char_count": len(chunk_text),
                        "format": "text",
                    },
                }
            )
            chunk_idx += 1
            current_chunk = []
            current_length = 0

        current_chunk.append(para)
        current_length += len(para)

    # Flush remaining
    if current_chunk:
        chunk_text = "\n\n".join(current_chunk)
        records.append(
            {
                "text": chunk_text,
                "source": filename,
                "metadata": {
                    "chunk": chunk_idx,
                    "char_count": len(chunk_text),
                    "format": "text",
                },
            }
        )

    logger.info(
        "[Parser] Text '%s': %d chars → %d records",
        filename,
        len(content),
        len(records),
    )
    return records


def _clean_pdf_text(text: str) -> str:
    """Clean common PDF extraction artifacts."""
    # Fix hyphenated line breaks
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove page headers/footers (lines that are just numbers)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    return text.strip()
